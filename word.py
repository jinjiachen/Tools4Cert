#!/bin/python
#coding:utf-8
"""
Author: Michael Jin
Date: 2022-08-11
"""


import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from excel import get_UC,get_data
from PyPDF2 import PdfFileMerger,PdfWriter,PdfReader,PdfMerger
import PyPDF2
import time
if os.name=='nt':
    from win32com.client import Dispatch
    import xlwings as xw


def Menu():
    choice=input('请输入你的选择：\n1.生成年检报告\n2.提取数据\n3.doc转docx\n4.批量doc转PDF\n5.合并pdf\n6.doc转pdf\n7.pdf加水印\ndft:生成草稿报告\ninit:初始化年检报告\nat(auto rotate):自动翻转PDF文件\npr(pdf replacement):pdf文件替换指定页\npe(pdf extract):提取部分pdf页数\npm(pdf merge):合并两个PDF文件')
    if choice=='uc':
        path_xls=input('请输入需要做年检的报告（excel)的文件夹路径')
#        path_doc=input('请输入年检报告(word)的路径')
        path_doc=r'J:\Tools4Cert\template\SFT-ETL-OP-29a Unlisted Component Acceptance Report_200331.docx'
        project=input('请输入项目号：')
        control_No=input('请输入控制号：')
        sample=input('请输入样品号：')
        path_doc=path_doc.replace('"','')#去除"号,做预处理
        app=xw.App(visible=False,add_book=False)#创建app对象，传入Annual_checks函数
        path_doc=Annual_init(app,path_xls,path_doc,project,control_No,sample)
        for component in ['compressor','motor','smps','transformer','pwb','power unit','switch power supply']:
            Annual_checks_HA(app,path_xls,path_doc,component)
        app.kill()#关闭进程
    elif choice=='ucgt':
        path_xls=input('请输入需要做年检的报告（excel)的文件夹路径')
#        path_doc=input('请输入年检报告(word)的路径')
        path_doc=r'J:\Tools4Cert\template\SFT-ETL-OP-29a Unlisted Component Acceptance Report_200331.docx'
        project=input('请输入项目号：')
        control_No=input('请输入控制号：')
        sample=input('请输入样品号：')
        path_doc=path_doc.replace('"','')#去除"号,做预处理
        app=xw.App(visible=False,add_book=False)#创建app对象，传入Annual_checks函数
        path_doc=Annual_init(app,path_xls,path_doc,project,control_No,sample)
        for component in ['compressor','motor','smps','transformer','pwb','power unit']:
            Annual_checks_GT(app,path_xls,path_doc,component)
        app.kill()#关闭进程
    elif choice=='2':
        file_type=input('提取的数据源类型：\n1.word 2.excel')
        if file_type=='1':
            data=input('请输入要提取的数据文件的路径：')
            data=data.replace('"','')
            docx=Document(data)
            new_docx=Document()
            tables=docx.tables
            print('the numbers of tables:',len(tables))
            table_components=find_table(tables,'hello')
            print('找到部件清单的表格：',table_components)
            rows_value=get_more(table_components)
            print('获取的数据行数：',len(rows_value))
            new_table=new_docx.add_table(rows=len(rows_value),cols=6,style="Table Grid")
            print('生成的新表格的行数：',len(new_table.rows))
            write_table(new_table,rows_value)
            new_docx.save(data[:-5]+'output.docx')
        elif file_type=='2':
            data=input('请输入要提取的数据文件的路径：')
            data=data.replace('"','')
            data_start=int(input("Please input the start line of data:"))
            data_end=int(input("Please input the end line of data:"))
            col1=input("Please choose four columns of data (1/4):")
            col2=input("Please choose four columns of data (2/4):")
            col3=input("Please choose four columns of data(3/4):")
            col4=input("Please choose four columns of data (4/4):")
            col5=input("是否有单独提供证书，如有，请指出证书号所在列")
            app=xw.App(visible=True,add_book=False)
            app.display_alerts=False #取消警告
            app.screen_updating=False#取消屏幕刷新
            wb_data=app.books.open(data)
            for sheet in wb_data.sheets:
                print(sheet)
                if sheet.name=='4.0 Components':
                    print('find',sheet.name)
                    sht_data=sheet
                    break
    #                print(sht_data.name)
                else:
                    sht_data=wb_data.sheets[0]
            source_data=get_data(sht_data,data_start,data_end,col1,col2,col3,col4,col5)#提取到的原始数据
            print(source_data)
            wb_data.close()
            app.kill()
            #开始将数据写入word
            new_docx=Document()
            new_table=new_docx.add_table(rows=len(source_data),cols=6,style="Table Grid")
            print('生成的新表格的行数：',len(new_table.rows))
            write_table(new_table,source_data)
            new_docx.save(data[:-5]+'output.docx')
    elif choice=='3':
        path=input('请输入需要转换的doc文件路径：')
        path=path.replace('"','')
        doc2docx(path)
    elif choice=='4':
        path=input('请输入需要转换的doc/docx文件夹路径：')
        path=path.replace('"','')
        docs2pdfs(path)
    elif choice=='5':
        target_path = input('PDF的文件夹路径:')
        pdf_merge(target_path)
    elif choice=='6':
        path=input('请输入需要转换的doc/docx文件路径：')
        path=path.replace('"','')
        doc2pdf(path)
    elif choice=='7':
        pdf_file=input('请输入pdf文件的路径:')
#        pdfWriter = PyPDF2.PdfFileWriter()      # 用于写pdf
#        pdfReader = PyPDF2.PdfFileReader(pdf_file)   # 读取pdf内容
        watermark='K:\Database\watermark.pdf'
        add_watermark(pdf_file,watermark)
#        # 遍历pdf的每一页,添加水印
#        for page in range(pdfReader.numPages):
#            page_pdf = add_watermark(watermark, pdfReader.getPage(page))
#            pdfWriter.addPage(page_pdf)
#        
#        with open(pdf_file, 'wb') as target_file:
#            pdfWriter.write(target_file)
    elif choice=='dft':
        path=input('请输入需要转换的doc文件路径：')
#        filename=os.path.basename(path)
#        dirname=os.path.dirname(path)
#        print(filename)
#        print(dirname)
        watermark='K:\Database\watermark.pdf'
        path=path.replace('"','')
        new_pdf=doc2pdf(path)
        print(new_pdf)
#        add_watermark(path[:-3]+'pdf',watermark)
        add_watermark(new_pdf,watermark)
        os.remove(new_pdf)
    elif choice=='init':
        pass
    elif choice=='at':
        pdf_path=input('请输入要翻转的PDF文件路径:')
        action=input('请输入需要旋转的方向：\nleft:向左旋转90\nright:向右旋转90')
        page_rotation(pdf_path,action)
    elif choice=='pr':
        input_pdf1 = input('Please input the first pdf')
        input_pdf1=input_pdf1.replace('"','')
        input_pdf2 = input('Please input the second pdf')
        input_pdf2=input_pdf2.replace('"','')
        page_to_replace=input('page to be replaced')
        replacement_page=input('page to replace')
        
        # 获取并验证页码参数
        try:
            page_to_replace = int(page_to_replace)
            replacement_page = int(replacement_page)
        except ValueError:
            print("错误: 页码必须是整数")

        if not replace_page(input_pdf1, input_pdf2, page_to_replace, replacement_page):
            sys.exit(1)
    elif choice=='pe':
        pdf=input('please input the pdf file')
        pdf=pdf.replace('"','')
        output_pdf=pdf[:-4]+'_output.pdf'
        pages=input('Please input the page range')
        pages=selection_sort(pages)
        extract_pages(pdf,output_pdf,pages)
    elif choice=='pm':
        pdf1=input('please input the first pdf file')
        pdf1=pdf1.replace('"','')
        pdf2=input('please input the second pdf file')
        pdf2=pdf2.replace('"','')
        output_pdf=pdf1[:-4]+'_output.pdf'
        merge_2pdfs(pdf1,pdf2,output_pdf)    




def add_watermark(pdf_path,watermark):
    """
    将水印pdf与pdf的一页进行合并
    :param water_file:
    :param page_pdf:
    :return:
    """
    pdf_water = PyPDF2.PdfFileReader(watermark)
    pdf_file = PyPDF2.PdfFileReader(pdf_path)   # 读取pdf内容
    pdfWriter = PyPDF2.PdfFileWriter()      # 用于写pdf
#    pdf_file.mergePage(pdf_file.getPage(0))
    # 遍历pdf的每一页,添加水印
    for page in range(pdf_file.numPages):
        pdf_page=pdf_file.getPage(page)
        pdf_page.mergePage(pdf_water.getPage(0))
        pdfWriter.addPage(pdf_page)

#    print(os.path.dirname(pdf_path))
    
    ###以新文件保存
    with open(pdf_path[:-4]+'_draft.pdf', 'wb') as target_file:
        pdfWriter.write(target_file)
    return target_file


def doc2docx(path):#将doc文件转换成docx
    '''path具体到文件'''
    w = Dispatch('Word.Application') #打开word程序
    w.Visible = 0 #后台运行，不可见
    w.DisplayAlerts = 0
    doc = w.Documents.Open(path) #打开对应的doc文件
    newpath = os.path.splitext(path)[0] + '.docx' #生成docx文件路径
    doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)#12为docx的类型
    doc.Close()
    w.Quit()
#    os.remove(path)#不删除源文件
    return newpath


def doc2pdf(path):#将doc/docx文件转换成pdf
    '''path具体到文件'''
    w = Dispatch('Word.Application') #打开word程序
    w.Visible = 0 #后台运行，不可见
    w.DisplayAlerts = 0
    doc = w.Documents.Open(path) #打开对应的doc文件
    newpath = os.path.splitext(path)[0] + '.pdf' #生成docx文件路径
    doc.SaveAs(newpath, 17)#17为pdf的类型
    doc.Close()
    w.Quit()
#    os.remove(path)#不删除源文件
    return newpath

def docs2pdfs(path):#批量转换doc/docx为pdf文件
    '''输入文件夹路径即可'''
    files=[f for f in os.listdir(path) if f.endswith('.doc') or f.endswith('.docx')]
    file_path=[os.path.join(path, filename) for filename in files]
    for file in file_path:
        print(f'正在处理{file}')
        doc2pdf(file)


def pdf_merge(path):
    pdf_lst = [f for f in os.listdir(path) if f.endswith('.pdf')]
    pdf_lst = [os.path.join(path, filename) for filename in pdf_lst]
    
    file_merger = PdfFileMerger()
    for pdf in pdf_lst:
        file_merger.append(pdf)     # 合并pdf文件
    
    file_merger.write(os.path.join(path,'Draft_report.pdf'))
    pass


def find_table(tables,search_string):#关键词查找对应表格
    table_components=[]
    for table in tables:#遍历每一个表格
        if table.cell(0,0).text=='24.1' and table.cell(0,1).text=='TABLE: Critical components information':#标准TRF中零部件清单表格
            table_components.append(table)
        elif table.cell(0,0).text==search_string:#自定义关键字来搜索表格
            table_components.append(table)
    return table_components


def Search_table(tables,search_string):
    table_components=[]
    for table in tables:
#        print(type(table))
        for row in table.rows:
#            print(type(row.cells))
#            print(len(row.cells))
            for cell in row.cells:
#                print(cell.text)
                if cell.text.strip()==search_string:
#                if search_sring in cell.text:
#                    print(table)
                    table_components.append(table)
#                    break
    return table_components

def get_rows(table):#输入一个表格，返回一个以列表为元素的列表，每一个列表就是一行
    rows_values=[]
    for row in table.rows:#遍历表格的每一行
        if is_italic_row(row)==True:
            continue
        values=[]
#        if row.cells[0].text=='Object / part':
#            print('删除object所在行')
#            continue
#        elif row.cells[0].text=='hello':
#            print('删除hello所在行')
#            continue
#        elif row.cells[0].text=='Object / part No.':
#            print('删除Object / part No.所在行')
#            continue
#        for i in range(1,7):#针对table24.1,会返回8个cell, 取其中的1-6个，头和尾是重复的，不知道原因
        for i in range(0,6):#针对标准的TRF中table24.1
            values.append(row.cells[i].text)#把一行的数据放在value这个列表中
#            print(values)
#            break
        rows_values.append(values)#把value添加到rows_values这个总列表中
#            print(rows_values)
    return rows_values

def get_more(tables):#输入多个表格，返回每行的数据
    rows_values=[]
    for table in tables:
        rows_values=rows_values+get_rows(table)#把每个表格的行数据拼接起来
    return rows_values

def write_table(new_table,rows_value):#把获取的数据写入到新建的表格中
    for i in range(0,len(rows_value)-1):#遍历数据列表中的每一个元素
        print(f'写入数据第{i}行',rows_value[i])
#        for j in range(0,6):
        for j in range(0,len(rows_value[0])):
            new_table.rows[i].cells[j].text=rows_value[i][j]
            new_table.rows[i].cells[j].paragraphs[0].runs[0].font.name='Arial'
            new_table.rows[i].cells[j].paragraphs[0].runs[0].font.size=Pt(10)
            print(f'写入数据第{i}行第{j}列数据')

def is_italic_cell(cell):#判断表格中的cell是否为斜体,如果是空的，返回None，如果为斜体返回Ture，其他返回False
    if cell.text=='':
        return None
    if cell.paragraphs[0].runs[0].italic==True:
        return True
    else:
        return False
    
def is_bold_cell(cell):#判断表格中的cell是否为加粗
    if cell.paragraphs[0].runs[0].bold==True:
        return True
    else:
        return False

def is_italic_row(row):
    for cell in row.cells:
        if is_italic_cell(cell)==None:
            continue
        elif is_italic_cell(cell)==True:
            continue
        else:
            return False
            break 
    return True

        
def Annual_check(docx,data,component):#对具体的年检信息进行写入
#    table=docx.tables[0]
    table_content=find_table(docx.tables,'Unlisted Component')
    table_test=find_table(docx.tables,'Model No.')
    table_construction=find_table(docx.tables,'Model')
    print('找到目录表格:',len(table_content))
    print('找到耐压测试表格:',len(table_test))
    print('找到结构表格:',len(table_construction))
    flag=1#正常写入年检信息返回0，默认1
    for uc in data['uc_info']:
        if component in uc['name'].lower():#只处理指定的部件
            flag=0#检测到年检信息，更改返回值为0
            row_cells=table_content[0].add_row().cells#找到目录表格后增加一行，写入对应的数据
            row_cells[0].text=uc['name']
            row_cells[1].text=uc['manufacturer']
            row_cells[2].text=uc['model']
            row_cells[3].text=data['basic_info']['report']
            row_cells[4].text=str(uc['photo_no'])
            row_cells[5].text=str(uc['item_no'])


            #以下部分为多个绕组信息处理
            k=1
            while f'designation_{k}' in list(uc.keys()):
                row_cells=table_construction[0].add_row().cells#找到结构表格后增加一行，写入对应数据
                row_cells[0].text=uc['model']
                row_cells[1].text=uc[f'designation_{k}']
                row_cells[2].text=str(uc[f'wire_size_{k}'])
                row_cells[3].text=str(uc[f'resistance_{k}'])
#                row_cells[4].text=
                row_cells[5].text='Pass'
                k=k+1

            #当数据大于一行时，进行合并操作
            if k>2:
                total_rows=len(table_construction[0].rows)-1#获取结构表格总行数，-1是为了匹配索引
                col_0=table_construction[0].columns[0].cells#获取结构表格第一列的单元格

                #清除需要合并单元格中除第一格外的其他数据
                for col in range(total_rows-(k-3),total_rows+1):
                    print(f'清除结构表格第{col}行数据',col_0[col].text)
                    col_0[col].text=""

#                print('k:',k)
                #合并单元格
                print(f'合并结构表格第{total_rows-k+2}:{total_rows}行的单元格')
                table_construction[0].cell(total_rows-k+2,0).merge(table_construction[0].cell(total_rows,0))


            j=1
            while f'location_{j}' in list(uc.keys()):#找到测试表格后增加一行，写入对应数据
                row_cells=table_test[0].add_row().cells
                row_cells[0].text=uc['model']
#                row_cells[1].text=uc['manufacturer']
                row_cells[1].text=uc[f'location_{j}']
                row_cells[2].text=uc[f'rating'].split(',')[0]
                if type(uc[f'voltage_{j}'])==float:#浮点转化成整数的字符串
                    row_cells[3].text=str(int(uc[f'voltage_{j}']))#转化为字符串以防出错
                else:
                    row_cells[3].text=str(uc[f'voltage_{j}'])#转化为字符串以防出错
                row_cells[4].text=str(uc[f'time_{j}'])#转化为str，防止出错
                row_cells[5].text='Pass'
                j=j+1

            #当数据大于一行时，进行合并操作
            if j>2:
                total_rows=len(table_test[0].rows)-1#获取测试表格总行数，-1是为了匹配索引
                col_0=table_test[0].columns[0].cells#获取测试表格第一列的单元格
                col_1=table_test[0].columns[1].cells#获取测试表格第二列的单元格

                #清除需要合并单元格中除第一格外的其他数据
                for col in range(total_rows-(j-3),total_rows+1):
                    print(f'清除测试表格第{col}行数据',col_0[col].text)
                    col_0[col].text=""
#                    col_1[col].text=""

                print('j:',j)
                #合并单元格
                print(f'合并测试表格第{total_rows-j+2}:{total_rows}行的单元格')
                table_test[0].cell(total_rows-j+2,0).merge(table_test[0].cell(total_rows,0))
#                table_test[0].cell(total_rows-j+2,1).merge(table_test[0].cell(total_rows,1))
    return flag


def exit_file(file_path):#判断一个文件是否存在
    dirname=os.path.dirname(file_path)
    filename=os.path.basename(file_path)
    for file in os.listdir(dirname):
        if filename==file:
            return True

            
def Annual_checks(app,path_xls,path_doc,component):#对多个报告生成对应部件的年检报告
    files=[f for f in os.listdir(path_xls) if f.endswith('.xls') or f.endswith('.xlsm')] #列出所有的xls文件
    file_path=[os.path.join(path_xls, filename) for filename in files]#所有xls文件的绝对路径
    new_file=path_doc[:-4]+component+'.docx'#构造新的docx文件路径，即输出对应部件的年检报告
    print(file_path)
    for file in file_path:#遍历所有的xls文件，即所有需要做年检的报告
        print(f'正在处理{file}')
        wb=app.books.open(file)
        data=get_UC(wb)#提取相应的UC信息
        print(data)
        wb.close()
        if exit_file(new_file):
            docx=Document(new_file)#如果存在对应的年检报告，则在对应报告中添加
        else:
            docx=Document(path_doc)
        if not Annual_check(docx,data,component):#如果返回为0，则为正常写入数据，此时保存年检报告
            docx.save(new_file)

###按照GT的特殊需求，每一份单独的报告出一份年检报告
def Annual_checks_GT(app,path_xls,path_doc,component):#对多个报告生成对应部件的年检报告
    files=[f for f in os.listdir(path_xls) if f.endswith('.xls') or f.endswith('.xlsm')] #列出所有的xls文件
    file_path=[os.path.join(path_xls, filename) for filename in files]#所有xls文件的绝对路径
    print("\n".join(file_path))
    for file in file_path:#遍历所有的xls文件，即所有需要做年检的报告
        filename=os.path.basename(file)#文件名
        if filename.endswith('.xls'):
            folder=os.path.join(os.path.dirname(file),filename[:-4])#构造文件夹路径
        elif filename.endswith('.xlsm'):
            folder=os.path.join(os.path.dirname(file),filename[:-5])#构造文件夹路径
        os.system(f'md "{folder}"')#创建文件夹
        new_file=os.path.join(folder,os.path.basename(path_doc)[:-5])+f'{component}.docx'#构建年检报告的路径
        print(f'正在处理{file}')
        wb=app.books.open(file)
        data=get_UC(wb)#提取相应的UC信息
        print(data)
        wb.close()
        if exit_file(new_file):
            docx=Document(new_file)#如果存在对应的年检报告，则在对应报告中添加
        else:
            docx=Document(path_doc)
        if not Annual_check(docx,data,component):#如果返回为0，则为正常写入数据，此时保存年检报告
            docx.save(new_file)
            
###按照Molly的新要求，每一份单独的报告出一份年检报告,不同部件放在一起
def Annual_checks_HA(app,path_xls,path_doc,component):#对多个报告生成对应的年检报告
    files=[f for f in os.listdir(path_xls) if f.endswith('.xls') or f.endswith('.xlsm')] #列出所有的xls文件
    file_path=[os.path.join(path_xls, filename) for filename in files]#所有xls文件的绝对路径
    print("\n".join(file_path))
    for file in file_path:#遍历所有的xls文件，即所有需要做年检的报告
        filename=os.path.basename(file)#文件名
        if filename.endswith('.xls'):
            folder=os.path.join(os.path.dirname(file),filename[:-4])#构造文件夹路径
        elif filename.endswith('.xlsm'):
            folder=os.path.join(os.path.dirname(file),filename[:-5])#构造文件夹路径
        os.system(f'md "{folder}"')#创建文件夹
        print(f'正在处理{file}')

        #打开对应的报告提取信息
        wb=app.books.open(file)
        data=get_UC(wb)#提取相应的UC信息
        print(data)
        wb.close()

        new_file=os.path.join(folder,os.path.basename(path_doc)[:-5])+'.docx'#构建年检报告的路径
        if exit_file(new_file):
            docx=Document(new_file)#如果存在对应的年检报告，则在对应报告中添加
        else:
            docx=Document(path_doc)
        if not Annual_check(docx,data,component):#如果返回为0，则为正常写入数据，此时保存年检报告
            docx.save(new_file)


###修改TRF中的table 24.1
def update_components():#更新修改table24.1
    pass

###替换word中对应文字
def content_replace(documents,old_word,new_word,ptf='NO'):#替换对应文字，保持格式不变
    paragraphs=documents.paragraphs
    tables=documents.tables
    sections=documents.sections

    #查找所有的段落
    for paragraph in paragraphs:
        if old_word in paragraph.text:
            text=paragraph.text
            name=paragraph.runs[0].font.name
            size=paragraph.runs[0].font.size
            color=paragraph.runs[0].font.color.rgb
#            print(text,name,size,color)
            if ptf=='YES':
                print(f'找到{old_word},正在替换')
            update_text=text.replace(old_word,new_word)
            paragraph.text=update_text
            paragraph.runs[0].font.name=str(name)
            paragraph.runs[0].font.size=int(size)
#            paragraph.runs[0].font.color=str(color)

    #查找所有的表格
    for table in tables:
        for cell in table._cells:
            if old_word in cell.text:
                if ptf=='YES':
                    print(f'找到{old_word},正在替换')
                text=cell.text
                cell.text=text.replace(old_word,new_word)

    #查找所有的页眉
    for section in sections:
        for cell in section.header.tables[0]._cells:
            if old_word in cell.text:
                if ptf=='YES':
                    print(f'找到{old_word},正在替换')
                text=cell.text
                cell.text=text.replace(old_word,new_word)


###年检模板初始化
def Annual_init(app,path_xls,path_doc,project,control_No,sample):
    '''
    app:xlwings的实例
    path_doc(str):模板文件的路径
    project(str):年检项目号
    control_No(str):控制号
    sample(str):样品编号
    '''
    report_No=project+'-001'
    product='xxx'
    standard='xxx'

    docx=Document(path_doc)
    files=[f for f in os.listdir(path_xls) if f.endswith('.xls') or f.endswith('.xlsm')] #列出所有的xls文件
    file_path=[os.path.join(path_xls, filename) for filename in files]#所有xls文件的绝对路径
    #选一份报告获取基本的信息
    wb=app.books.open(file_path[0])#只选取一个报告
    data=get_UC(wb)#提取相应的UC信息
    wb.close()
    client_name=data['basic_info']['applicant']
    client_address=data['basic_info']['address']+', '+data['basic_info']['country']
    client_contact=data['basic_info']['contact']
    #一系列替换操作来初始化
    content_replace(docx,'CUSTOMER NAME',client_name)
    content_replace(docx,'<Client Name>',client_name)
    content_replace(docx,'<report no.>',report_No)
    content_replace(docx,'<issue_date>',time.strftime('%d-%m-%Y'))
    content_replace(docx,'<Control Number>',control_No)
    content_replace(docx,'<date>',time.strftime('%Y-%m-%d'))
    content_replace(docx,'<customer>',client_name)
    content_replace(docx,'<project>',project)
    content_replace(docx,'<sample no>',sample)
    content_replace(docx,'<product>',product)
    content_replace(docx,'<standard>',standard)
    content_replace(docx,'<Client Contact>',client_contact)
    content_replace(docx,'<Client Address>',client_address)
    #生成初始化文件并返回路径
    new_path=path_xls+'\\'+f'{project}.docx'
    docx.save(path_xls+'\\'+f'{project}.docx')
    return new_path

###PDF文件自动翻转
def page_rotation(old_file,action):
    """
    PDF页面旋转
    :param old_file: 需要旋转的PDF文件
    action(str):旋转方向
    :return:
    """
    print(f'PDF路径为：{old_file}')
    pdf = PdfReader(old_file)
    page_num = len(pdf.pages)
    pdf_writer = PdfWriter()
    for i in range(page_num):
        if action=='auto':
            size = pdf.pages[i].mediabox  # 获取页面大小值（长、宽）
            print(size)
    #        x, y = size.upper_right(), size.UpperRight_y()
            x=size[2]
            y=size[3]
            if x > y:
                # 顺时针旋转90度  90的倍数
                page = pdf.pages[i].rotate(90)
                # 逆时针旋转90度  90的倍数
                # page = pdf.getPage(i).rotateCounterClockwise(90)
                pdf_writer.add_page(page)
            else:
                # 不旋转
                page = pdf.pages[i].rotate(0)
                pdf_writer.add_page(page)
        elif action=='left':
            # 逆时针旋转90度  90的倍数
            page = pdf.pages[i].rotate(-90)
            pdf_writer.add_page(page)
        elif action=='right':
            # 顺时针旋转90度  90的倍数
            page = pdf.pages[i].rotate(90)
            pdf_writer.add_page(page)
    with open(old_file[:-4]+f'_{action}.pdf', 'wb') as f:
        print(f'输出PDF为：{f}')
        pdf_writer.write(f)

###用第二个PDF的某一页替换第一个PDF的某一页
def replace_page(input_pdf1, input_pdf2, page_to_replace, replacement_page):
    """
    input_pdf1 (str): 第一个PDF文件路径
    input_pdf2 (str): 第二个PDF文件路径
    page_to_replace (int): 第一个PDF中要替换的页码（从1开始）
    replacement_page (int): 第二个PDF中用于替换的页码（从1开始）
    """
    try:
        #格式预处理
        input_pdf1=input_pdf1.replace('"','')
        input_pdf2=input_pdf2.replace('"','')
        output_pdf=input_pdf1[:-4]+'_output.pdf'

        # 读取两个输入PDF
        reader1 = PdfReader(input_pdf1)
        reader2 = PdfReader(input_pdf2)
        writer = PdfWriter()
        
        # 检查页码是否有效
        total_pages1 = len(reader1.pages)
        total_pages2 = len(reader2.pages)
        
        if page_to_replace < 1 or page_to_replace > total_pages1:
            print(f"错误: 第一个PDF的页码 {page_to_replace} 超出范围 (1-{total_pages1})")
            return False
            
        if replacement_page < 1 or replacement_page > total_pages2:
            print(f"错误: 第二个PDF的页码 {replacement_page} 超出范围 (1-{total_pages2})")
            return False
        
        # 添加第一PDF的页面，替换指定页
        for i in range(total_pages1):
            if i == page_to_replace - 1:  # 找到要替换的页码
                writer.add_page(reader2.pages[replacement_page - 1])
            else:
                writer.add_page(reader1.pages[i])
        
        # 写入输出PDF
        with open(output_pdf, 'wb') as output_file:
            writer.write(output_file)
        
        print(f"成功用 {input_pdf2} 的第 {replacement_page} 页替换 {input_pdf1} 的第 {page_to_replace} 页")
        print(f"结果已保存到 {output_pdf}")
        return True
        
    except FileNotFoundError:
        print(f"错误: 找不到文件 {input_pdf1} 或 {input_pdf2}")
        return False
    except Exception as e:
        print(f"发生未知错误: {e}")
        return False


###提取pdf的部分页数
def extract_pages(input_pdf, output_pdf, page_numbers):
    """
    从PDF中提取指定页码并保存为新的PDF
    
    参数:
    input_pdf (str): 输入PDF文件路径
    output_pdf (str): 输出PDF文件路径
    page_numbers (list): 要提取的页码列表（从1开始）
    """
    try:
        # 读取输入PDF
        reader = PdfReader(input_pdf)
        writer = PdfWriter()
        
        # 检查页码是否有效
        total_pages = len(reader.pages)
        invalid_pages = [p for p in page_numbers if p < 1 or p > total_pages]
        
        if invalid_pages:
            print(f"错误: 页码 {', '.join(map(str, invalid_pages))} 超出范围 (1-{total_pages})")
            return False
        
        # 添加指定页到输出PDF
        for page_num in page_numbers:
            writer.add_page(reader.pages[page_num - 1])  # 注意：PyPDF2中页码从0开始
        
        # 写入输出PDF
        with open(output_pdf, 'wb') as output_file:
            writer.write(output_file)
        
        print(f"成功提取 {len(page_numbers)} 页到 {output_pdf}")
        return True
        
    except FileNotFoundError:
        print(f"错误: 找不到文件 {input_pdf}")
        return False
    except Exception as e:
        print(f"发生未知错误: {e}")
        return False

###对选择进行拆分排序然后输出列表
def selection_sort(selection):
    '''
    selection(str):测试的选择
    '''
    #对选择字符串做处理
    selection_sort=[]
    prehandle=selection.split(",")
    for i in prehandle:
        if "-" in i:
            j=i.split("-")
            j[0]=int(j[0])
            j[1]=int(j[1])
            for k in range(j[0],j[1]+1):
                k=str(k)
                selection_sort.append(k)
        else:
            selection_sort.append(i)
    #转化为整数
    selection=list(map(int,selection_sort))
    #对整数排序
    selection.sort()
    return selection


###合并两个PDF文件
def merge_2pdfs(file1_path, file2_path, output_path):
    """
    file1_path (str): 第一个PDF文件的路径
    file2_path (str): 第二个PDF文件的路径
    output_path (str): 合并后PDF文件的输出路径
    """
    # 检查输入文件是否存在
    for file_path in [file1_path, file2_path]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
    
    # 检查输出目录是否存在，不存在则创建
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 创建合并器对象
    merger = PdfMerger()
    
    try:
        # 添加并合并PDF文件
        merger.append(file1_path)
        merger.append(file2_path)
        
        # 写入合并后的文件
        merger.write(output_path)
        print(f"PDF文件已成功合并到: {output_path}")
    except Exception as e:
        print(f"合并PDF时出错: {e}")
    finally:
        # 关闭合并器对象
        merger.close()

if __name__=='__main__':
    while True:
        if os.name=='nt':
            os.system('cls')
        Menu()
        input('按任意键继续!')
