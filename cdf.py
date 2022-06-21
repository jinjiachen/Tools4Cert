#!/bin/python
#coding:utf-8
"""
Author: Michael Jin
Date: 2021-06-08
"""


import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

'''
def Search_table(tables,search_string):
    for table in tables:
#        print(table)
        for cells in table._cells:
            print(cells[0].text)
#            print(table._cells[0].text)
            for i in range(0,10):
                if cells[i].text == search_string:
                    print("Yes")
'''
####################################### main program #############################################
while True:
    try:
        print("This TDS is for IEC 60335-2-40:2018 in conjunction with IEC 60335-1:2010+A1:2013+A2:2016!")
        print("================================================== Program begin ==================================================")
#        job=input("Please input the project number:")
        document = Document('.\cr.docx') #Open the template document
        tables = document.tables #��ȡ�ĵ��б������б�
#        Search_table(tables,'24.1')
        table29 = tables[29] #��ȡ��30�����, table 24.1
        rows=table29.rows #��ȡ�����������
        columns=table29.columns #��ȡ�����������
#        print(rows[0].cells[1].paragraphs[0].runs[0].bold) #������������Ƿ�Ӵ�
        for i in range(2,len(rows)):
            print(rows[i].cells)
        if rows[0].cells[0].text=='24.1': #�жϱ�������ֵ�����
            print("yes")
#        for i in range(0,10):
#            table29.add_row() #����һ��
        ###����һ��ֱ��ͨ����Ԫ������ȡ����
#        cells = table5._cells #��ȡ��Ӧ������еĵ�Ԫ��
#        j=1
#        for i in range(3,5863,4):
#            cells[i].text = str(j)
#            j = j+1

        ###���������Ȼ�ȡ�ж����ٻ�ȡ�еĵ�Ԫ��
#        cols = table5.columns #��ȡ����е��ж���
#        col3 = cols[3] #��ȡ������
#        for i in range(0,1466):
#            col3.cells[i].text = "P"

        document.save('cdf.docx')

        print("==================================================Program END ==================================================") 
        flag=input("Press Enter to continue! Others to EXIT!")
        if flag!="":
            break
        else:
            os.system("cls")

    except:
        print("==================================================Program END ==================================================") 
        print("Error, please contact Michael.jc.jin@intertek.com")
        break

#document.save('D:\Downloads\Tools4Cert-master\demo2.docx')
