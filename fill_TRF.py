#!/bin/python
#coding:utf-8
"""
Author: Michael Jin
Date: 2021-06-08
"""


import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from word import find_table

def find_table(tables,search_string):#关键词查找对应表格
    table_components=[]
    for table in tables:#遍历每一个表格
        if table.cell(0,1).text==search_string:#自定义关键字来搜索表格
            table_components.append(table)
    return table_components


###用于把TRF文件条款部分自动填充序号
def fill_number(document,method):
#    '''
#    document: Document的实例
#    method:cells和rows两种方法，cells速度快，但是不够直观，rows直观，概念清晰，但是速度慢
#    '''

    tables = document.tables #获取文档中表格对象列表
    table5 = tables[5]#条款判断至annex s结束
    table7 = tables[7] #条款判断部分ANNEX S开始的表
    if method=='cells':
        cells = table5._cells #判断部分一共分两个表，这是第一个表的序号处理
        total_rows=len(table5.rows)
        j=1
        for i in range(3,total_rows*4+1,4):
            cells[i].text = str(j)
            j = j+1

        cells = table7._cells #判断部分一共分两个表，这是第二个表的序号处理
        total_rows=len(table7.rows)
        j=1
        for i in range(3,total_rows*4+1,4):
            cells[i].text = str(j)
            j = j+1
    elif method=='rows':
        ###判断部分一共分两个表，这是第一个表的序号处理
        j=1
        total_rows=len(table5.rows)
        for row in range(0,total_rows):
            print(f'{row}/{total_rows}')
            cells=table5.row_cells(row)
            cells[3].text=str(j)
            j=j+1
    
        ###判断部分一共分两个表，这是第二个表的序号处理
        j=1
        total_rows=len(table7.rows)
        for row in range(0,total_rows):
            print(f'{row}/{total_rows}')
            cells=table7.row_cells(row)
            cells[3].text=str(j)
            j=j+1

    document.save(TRF_path[:-5]+'_filled.docx')


###自动填充不适用的章节
def auto_NA(table,start,end):
    '''
    table:条款分两张表，指定是哪一张表格
    start(int):起始行
    end(int):结尾行
    '''
    cells = table._cells #获取所有的单元格
    for row in range(start,end+1):
        if cells[row*4-2].text=='':
            cells[row*4-1].text='N/A'#行数和单元格之间的对应关系，cell=4*row-1
        elif cells[row*4-2].text!='':
            cells[row*4-1].text='--'#行数和单元格之间的对应关系，cell=4*row-1

###某一行的判定
def fill_line(table,row,value='N/A'):
    '''
    table:表格对象
    row(int):行号
    value(str):P or N/A
    '''
    cells = table._cells #获取所有的单元格
    cells[4*row-1].text=value
    

###某一行的注释
def comment(table,row,value):
    '''
    table:表格对象
    row(int):行号
    value(str):想写的内容
    '''
    cells = table._cells #获取所有的单元格
    cells[4*row-2].text=value
    


###条款判断,适用于2-40v
def verdict(document):
    '''
    document:文档对象
    '''
    ###初始化
    tables=document.tables
    table5=tables[5]
    table7=tables[7]


    ###参数
    volt=input('请输入器具的额定电压：')
    phase=input('是否为单相器具：')
#    if phase=='n':
#        phase_number=input('多相器具请输入相数：')
    acdc=input('ac or dc?')
    if acdc=='ac':
        freq=input('请输入频率:')
    power=input('请输入额定功率:')
    current=input('请输入额定电流：')
    appliance_type=input('请选择器具类型：')
    AC_type=input('分体还是移动')
    grade=input('器具为I,II,III?')
    use_env=input('器具的使用环境')
    laundry=input('是否在洗衣房使用')
    public=input('大众是否可触及')
    flammable=input('是否可燃')
    if flammable=='y':
        leak=input('是否做泄露测试(Annex FF)')
        detection=input('是否采用可燃冷媒探测系统')
        arrest=input('采用火焰捕捉外壳(Annex NN)?')
    else:
        leak='n'
        detection='n'
        arrest='n'
    UV=input('是否包含UV-C')



    ###cl.5
    fill_line(table5,1,'--')
    fill_line(table5,2,'P')


    ###cl.5.2
    fill_line(table5,3,'P')
    if flammable=='y':
        if leak=='y':
            fill_line(table5,4,'P')
        if detection=='y':
            fill_line(table5,5,'P')
        else:
            fill_line(table5,5,'N/A')
        if arrest=='y':
            fill_line(table5,6,'P')
        else:
            fill_line(table5,6,'N/A')
    else:
        auto_NA(table5,4,6)


    ###cl.5.6
    fill_line(table5,7,'P')

    
    ###cl.5.7
    fill_line(table5,8,'P')


    ###cl.5.10
    if AC_type=='split air conditioner':
        fill_line(table5,9,'P')
        fill_line(table5,10,'P')
        fill_line(table5,11,'P')
    else:
        auto_NA(table5,9,11)


    ###cl.6
    fill_line(table5,12,'--')
    

    ###cl.6.1
    fill_line(table5,13,'P')
    comment(table5,13,grade)


    ###cl.6.2
    fill_line(table5,4,'--')
    if use_env=='indoor use only':
        fill_line(table5,15,'N/A')
        fill_line(table5,16,'P')
    elif use_env=='outdoor use':
        fill_line(table5,15,'P')
        fill_line(table5,16,'N/A')
    if laundry=='y':
        fill_line(table5,17,'P')
    else:
        fill_line(table5,17,'N/A')


    ###cl.6.101
    if public=='y':
        fill_line(table5,18,'P')
        comment(table5,18,'accessible to general public')
    elif public=='n':
        fill_line(table5,18,'P')
        comment(table5,18,'not accessible to general public')


    ###cl.7
    fill_line(table5,19,'--')


    ###cl.7.1
    fill_line(table5,20,'P')
    comment(table5,20,volt)
    if phase=='y':
        fill_line(table5,21,'N/A')
    elif phase=='n':
        fill_line(table5,21,'P')
        comment(table5,21,'refer to the marking label')
    if acdc=='ac':
        fill_line(table5,22,'P')
        comment(table5,22,freq)
    elif acdc=='dc':
        fill_line(table5,22,'N/A')
    if power!='':
        fill_line(table5,23,'P')
        comment(table5,23,power)
    else:
        fill_line(table5,23,'N/A')
    if current!='':
        fill_line(table5,24,'P')
        comment(table5,24,current)
    else:
        fill_line(table5,24,'N/A')
    fill_line(table5,25,'P')
    comment(table5,25,'refer to the marking label')
    fill_line(table5,26,'P')
    comment(table5,26,'refer to the marking label')













    ###Annex B
    if appliance_type!='battery-operated appliance':
        auto_NA(table5,1296,1344)

    ###Annex C
    auto_NA(table5,1345,1347)


    ###Annex F
    auto_NA(table5,1361,1390)


    ###Annex H
    auto_NA(table5,1406,1432)


    ###Annex J
    auto_NA(table5,1433,1442)


    ###Annex K
    pass


    ###Annex M
    pass


    ###Annex N
    pass


    ###Annex R
    auto_NA(table5,1475,1520)



    ###Annex S
    if appliance_type!='battery-operated appliance':
        auto_NA(table7,1,42)


    ###Annex T
    if UV=='no':
        auto_NA(table7,43,62)


    ###Annex DD
    if flammable=='no':
        auto_NA(table7,63,234)


    ###Annex EE
    pass


    ###Annex FF
    if flammable=='no':
        auto_NA(table7,255,285)


    ###Annex GG
    if flammable=='no':
        auto_NA(table7,286,606)



    ###Annex JJ
    auto_NA(table7,607,615)



    ###Annex KK
    if flammable=='no':
        auto_NA(table7,616,648)



    ###Annex LL
    if flammable=='no':
        auto_NA(table7,649,706)


    ###Annex MM
    if flammable=='no':
        auto_NA(table7,707,717)


    ###Annex NN
    if flammable=='no':
        auto_NA(table7,718,724)


    ###Annex PP
    if flammable=='no':
        auto_NA(table7,725,747)


    ###Annex QQ
    if flammable=='no':
        auto_NA(table7,748,811)

    return document 



####################################### main program #############################################
if __name__=='__main__':
    TRF_path=input('Please input the TRF file:')
    document = Document(TRF_path) #Open the template document
    verdict(document)
    document.save(r'J:\Tools4Cert\TRF\test.docx')


#    while True:
#        try:
#            print("This TDS is for IEC 60335-2-40:2018 in conjunction with IEC 60335-1:2010+A1:2013+A2:2016!")
#            print("================================================== Program begin ==================================================")
#    #        job_no=input("Please input the project number:")
#            TRF_path=input('Please input the TRF file:')
#            document = Document(TRF_path) #Open the template document
##            fill_number(document,'cells')
#            verdict(document)
#    
#            document.save(r'J:\Tools4Cert\TRF\test.docx')
#    
#            print("==================================================Program END ==================================================") 
#            flag=input("Press Enter to continue! Others to EXIT!")
#            if flag!="":
#                break
#            else:
#                os.system("cls")
#    
#        except:
#            print("==================================================Program END ==================================================") 
#            print("Error, please contact Michael.jc.jin@intertek.com")
#            break
#    

