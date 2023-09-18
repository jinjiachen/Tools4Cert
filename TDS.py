#!/bin/python

import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#ClientName="xxxx"
#ProjectNo="00000"
#Appliance="xxx"
#ProjectEngineer="sss"
#TestEngineer="aaa"
#TestingDate="xxx"
#Reviewer="xxx"
#Models="xxx"
#Ratings="bbb"
#Standard="xxx"
#SampleNo="111111"
#document.add_picture('D:\Python27\Tools4Cert-master\34.jpg', width=Inches(1.25))
#run=body.add_run('test one') #insert the pattern
#font=run.font   
#font.name="Arial"   #set the font style
#font.size=Pt(10)    #set the font size
#section = document.sections[0] #choose the first section
#header = section.header #insert the first section's header
#footer = section.footer #insert the first section's footer
#para=header.paragraphs[0]
#para.text="intertek\t page __ of __\n"+"Client Name:"+ClientName+"\t"+"Porject No:"+ProjectNo+"\n"+"Project Engineer:"+ProjectEngineer+"\t"+"Reviewer:"+Reviewer+"\n"+"Test Engineer:"+TestEngineer+"\t"+"Testing Date:"+TestingDate
#para_foot=footer.paragraphs[0]
#para_foot.text="version 0.1"

#table=document.add_table(rows=5,cols=3) #insert table
#table.style='Light Shading Accent 1'
#cell=table.cell(0,1)
#cell.text="clause1"
verdict='P / F / NA'
values=[
        ("7.14","Marking Durability Test","15s with water, 15s with petroleum","--",verdict),
        ("22.11&22.34","Push, Pull and Torque Test","As specified in standard","EC2092, EC6621, EC2162, EC2426",verdict),
        ("8.1.1","Protection Against Access to Live Parts","As specified in standard","EC2162, EC3826",verdict),
        ("8.1.2","Protection Against Access to Live Parts","As specified in standard","EC3091",verdict),
        ("8.1.3","Protection Against Access to Live Parts","As specified in standard","EC2017",verdict),
        ("8.1.4","Protection Against Access to Live Parts","As specified in standard","EC3175, EC5800",verdict),
        ("8.1.5","Protection Against Access to Live Parts","As specified in standard","EC2162, EC3826",verdict),
        ("8.2","Protection Against Access to Live Parts","As specified in standard","EC2162, EC3826",verdict),
        ("10.1&10.2","Power input/Current Deviation","Refer to table below for details","EC6565, EC5865, EC5936",verdict),
        ("11.8","Heating Test","Refer to below for details","EC6565, EC5865, EC5936, EC3102, EC4232",verdict),
        ("13.2 & 13.3","Leakage Current Test & Electric Strength Test","Refer to table below for details","EC3175, EC3074, EC2834, EC6565, EC5132",verdict),
        ("14","Transient Overvoltage","Refer to table below for details","--",verdict),
        ("15.2","IP Test","As specified in standard","--",verdict),
        ("15.3","Overflow Test","As specified in standard","--",verdict),
        ("15.101","Spillage Test","As specified in standard","EC2969, EC2834, EC2385",verdict),
        ("16.2 & 16.3","Leakage Current Test & Electric Strength Test","Refer to talbe 1-5","EC2743, EC5800, EC5132, EC2834",verdict),
        ("17","Overload Protection Temperature Test","Refer to table below","EC5865, EC3102, EC5936",verdict),
        ("19.2&19.3","Abnormal Operation Restricted Heat Dissipation & Overload Test","As specified in standard","EC5936, EC5865, EC3102",verdict),
        ("19.4","Operation with any defect","As speficied in standard","EC5936, EC5865",verdict),
        ("19.5","Short-circuited the Sheath and N conductor","As specified in standard","--",verdict),
        ("19.6","Abnormal Operation-PTC","As specified in standard","EC5936, EC5865, EC3102",verdict),
        ("19.7","Locking Test for the Motor","Refer to table below","EC5132, EC3102, EC2834, EC5800, EC2743",verdict),
        ("19.8","Three phase motor","Refer to table below","EC5865, EC5936",verdict),
        ("19.10","Tests for Series Motors","As specified in standard","EC5865, EC5936",verdict),
        ("19.11.2","Fault Conditions of Electronic Circuit","Refer to table below","EC5865, EC5936",verdict),
        ("19.11.4.8","Voltage Drop Test","Refer to table below","EC5865, EC5936",verdict),
        ("19.12","Tests for Miniature Fuse-link","Refer to table below","EC5865, EC5936, EC5800",verdict),
        ("19.101","Restriction of Heat Transfer Medium Flow","As specified in standard","EC5865, EC5936, EC3102",verdict),
        ("19.102","Abnormal Temperature of indoor water","As specified in standard","EC5865, EC5936, EC3102",verdict),
        ("19.103","Abnormal Ambient Temperature","As specified in standard","EC5865, EC5936, EC3102, EC2605",verdict),
        ("19.104","Cover Test for Appliance with Supplementary Heaters","As specified in standard","EC5865, EC5936, EC3102",verdict),
        ("20.1","Stability Test","(   ) inclined","EC4261",verdict),
        ("20.2","Mechanical Hazard","As specified in standard","EC2162",verdict),
        ("21.1","Spring Hammer Test","As specified in standard","EC5553",verdict),
        ("Annex EE","Pressure Tests","Refer to table below","EC5768",verdict),
        ("21.2","Strength of Solid Insulation & Viberation Test","As specified in standard","--",verdict),
        ("22.3","Undue Strain Test on Socket-Outlet","As specified in standard","EC5076",verdict),
        ("22.5","Plug Discharge Test","Refer to table below","EC2567, EC3175, EC5132",verdict),
        ("22.6","Water Leakage Test","As specified in standard","EC2615",verdict),
#        ("22.11","Stability of Non-detachable Parts","Refer to table below","EC2092",verdict),
        ("22.12","Pull Test","(   )N","EC2092",verdict),
        ("22.16","Cord Reel Abrasion Test","As specified in standard","EC2285, EC2384",verdict),
        ("22.24","Bare Heating Elements","As specified in standard","--",verdict),
        ("22.32","Ageing Test of Rubber & Test of Ceramic Material","As specified in standard","--",verdict),
        ("22.42","Protective Impedance","As specified in standard","--",verdict),
        ("22.47","Water Mains Pressure Test","(   )MPa, 5mins, No Leakage","EC2468, EC3667",verdict),
        ("22.57","UV-C Radiation","As specified in standard","--",verdict),
        ("Annex T","UV-C Radiation","Refer to table below","--",verdict),
        ("22.104","Water Pressure Test for Containers","(   )MPa, 5mins, No Leakage","EC5768",verdict),
        ("22.108","Vacuum Pressure Impulses for Storage Tanks","(   )MPa, 15mins, No Deformation","EC5768",verdict),
        ("22.110","Operation of Non-self-resetting thermal cut-outs","As specified in standard","--",verdict),
        ("22.127-22.129","Irradiance Limit Test","As specified in standard","EC5912, EC4255",verdict),
        ("23.3","Internal Wiring Flexing Test","(\t) times for the flexing conductors","EC2285",verdict),
        ("23.5","Insulation of Internal Wiring Test","2000 V, 15 minutes","EC2834",verdict),
        ("24.5","Capacitor voltage","Rated voltage:\nMeasured voltage:","EC6081, EC4937, EC5132",verdict),
        ("23.101","Radiation resistance of internal wiring","After conditioning in Annex OO, 2000V, 15 minutes applied","EC2834",verdict),
        ("25.2","Electric strength for multiple supply","1250V, 60s","EC2834",verdict),
        ("25.14","Cord Flexing Test","As specified in standard","EC2667, EC2211",verdict),
        ("25.15","Power Cord Pull and Torque Test","Mass of appliance: (\t) kg\nPull force: (\t) N\nTorque: (\t) Nm\nMovement distance: (\t)mm","EC2092, EC6622, EC2843",verdict),
        ("26.5","Conductor Escape Test","As specified in standard","--",verdict),
        ("27.5","Ground Impedance Test","(\t)Ohm","EC4291",verdict),
        ("29","Creepage Ditance and Clearance","As specified in standard","EC2584, EC2843",verdict),
        ("30.1","Ball Pressure Test","Refer to table below","EC3304, EC2132",verdict),
        ("30.2.3 & 30.2.4","Glow Wire Test & Needle Flame Test","Refer to table below","EC2764, EC2072",verdict),
        ("31","Salt Mist Test","As specified in standard","--",verdict),
        ("32.101","UV-C Irradiance Test","Measured UV-C spectral irradiance: (\t)uW/cm2","--",verdict),
        ("Annex N","Proof Tracking Test","Refer to table below","EC2071",verdict),
#        ("Annex B","Appliance Powered by Rechargeable Batteries","--",verdict),
#        ("Annex D","Motor with Thermal Protectors","--",verdict),
#        ("Annex H","Switch Endurance Test","--",verdict),
#        ("Annex I","Motors having basic insulation that is inadequate for the rated voltage of the appliance","--",verdict),
#        ("Annex EE","Pressure Tests","--",verdict),
        ("Annex FF","Leakage Simulation Tests","Refer to table below","EC5382, EC6022, EC6023",verdict),
#        ("Annex GG","Refrigerant Charge","--",verdict)
        ]

Test_clauses=[
        ("10.1&10.2","Power input/Current Deviation","Refer to table below for details","EC6565, EC5865, EC5936","P/F/NA"),
        ("11.8","Heating Test","Refer to below for details","EC6565, EC5865, EC5936, EC3102, EC4232","P/F/NA"),
        ("13.2 & 13.3","Leakage Current Test & Electric Strength Test","Refer to table below for details","EC3175, EC3074, EC2834, EC6565, EC5132","P/F/NA"),
        ("16.2 & 16.3","Leakage Current Test & Electric Strength Test","Refer to talbe 1-5","EC2743, EC5800, EC5132, EC2834","P/F/NA"),
        ("17","Overload Protection Temperature Test","Refer to table below","EC5865, EC3102, EC5936","P/F/NA"),
        ("19.2&19.3","Abnormal Operation Restricted Heat Dissipation & Overload Test","As specified in standard","EC5936, EC5865, EC3102","P/F/NA"),
        ("19.7","Locking Test for the Motor","Refer to table below","EC5132, EC3102, EC2834, EC5800, EC2743","P/F/NA"),
        ("19.11.2","Fault Conditions of Electronic Circuit","Refer to table below","EC5865, EC5936","P/F/NA"),
        ("19.11.4.8","Voltage Drop Test","Refer to table below","EC5865, EC5936","P/F/NA"),
        ("19.101","Restriction of Heat Transfer Medium Flow","As specified in standard","EC5865, EC5936, EC3102","P/F/NA"),
        ("19.102","Abnormal Temperature of indoor water","As specified in standard","EC5865, EC5936, EC3102","P/F/NA"),
        ("19.103","Abnormal Ambient Temperature","As specified in standard","EC5865, EC5936, EC3102, EC2605","P/F/NA"),
        ("19.104","Cover Test for Appliance with Supplementary Heaters","As specified in standard","EC5865, EC5936, EC3102","P/F/NA"),
        ("Annex EE","Pressure Tests","Refer to table below","EC5768","P/F/NA"),
        ("Annex T","UV-C Radiation","Refer to table below","--","P/F/NA"),
        ("22.104","Water Pressure Test for Containers","(   )MPa, 5mins, No Leakage","EC5768","P/F/NA"),
        ("30.1","Ball Pressure Test","Refer to table below","EC3304, EC2132","P/F/NA"),
        ("30.2.3 & 30.2.4","Glow Wire Test & Needle Flame Test","Refer to table below","EC2764, EC2072","P/F/NA"),
        ("Annex N","Proof Tracking Test","Refer to table below","EC2071","P/F/NA"),
        ("Annex FF","Leakage Simulation Tests","Refer to table below","EC5382, EC6022, EC6023","P/F/NA"),
        ]

def Content(values):
    table=document.add_table(rows=1,cols=5,style="Table Grid")
    table.style.font.name="Arial"   #set the font of table to Arial style
    table.style.font.size=Pt(10)    #set the font size
#    table.autofit()
    row=table.rows[0].cells
    row[0].text="Clause"
#    row[0].text.name="Arial"	
    row[1].text="Test Description"
    row[2].text="Test Condition"
    row[3].text="Equipment Used"
    row[4].text="Verdit"
    for u,v,w,x,y in values:
        row_cells=table.add_row().cells     #add rows of table one by one
        row_cells[0].text=u
        row_cells[1].text=v
        row_cells[2].text=w
        row_cells[3].text=x
        row_cells[4].text=y
    document.add_page_break()

def Clause10():
#    table_head=document.add_table(rows=3,cols=2,style="Table_Normal") 
#    table_head.style.font.name="Arial"  #effect only when specify the table style	
#    table_head.style.font.size=Pt(10)   #effect only when specify the table style	
#    col=table_head.columns[0].cells
#    col[0].text="Test Item:"
#    col[1].text="Test Method:"
#    col[2].text="Test Result:"
    p=document.add_paragraph().add_run("Test Item:\tPower input/current deviation(Clause 10.1&10.2)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"
#    p1=document.add_paragraph()
#    p1.add_run("Table 1-1").font.name="Arial"
#    p1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER #set the alignment to center
    #add testing table
#    document.add_paragraph().add_run("Power Input Deviation:").font.name="Arial"
#    table_test=document.add_table(rows=8,cols=6,style="Table Grid")     ##Table for power input
#    values=[("Model","Rated Power Input(W)","Measured Power Input(W)","Deviation","Calculated Deviatoin","Mode")]
#    for row in table_test.rows:     #set the row height 
#        row.height=Pt(20)
#    row_cells=table_test.rows[0].cells  #choose the first row of table
#    for u,v,w,x,y,z in values:
#        row_cells[0].text=u
#        row_cells[1].text=v
#        row_cells[2].text=w
#        row_cells[3].text=x
#        row_cells[4].text=y
#        row_cells[5].text=z
#    document.add_paragraph()        ##blank line
#    document.add_paragraph().add_run("Current Deviation:").font.name="Arial"
#    table_test=document.add_table(rows=8,cols=6,style="Table Grid")     ##Table for rated current
#    values=[("Model","Rated Current(A)","Measured Current(A)","Deviation","Calculated Deviatoin","Mode")]
#    for row in table_test.rows:     #set the row height 
#        row.height=Pt(20)
#    row_cells=table_test.rows[0].cells  #choose the first row of table
#    for u,v,w,x,y,z in values:
#        row_cells[0].text=u
#        row_cells[1].text=v
#        row_cells[2].text=w
#        row_cells[3].text=x
#        row_cells[4].text=y
#        row_cells[5].text=z
    table_test=document.add_table(rows=8,cols=5,style="Table Grid")     ##Table for power input
    values=[("Model","Operating Mode","Test Condition (C)","Test Voltage (V)","Test Frequency (Hz)","Measured Power Input (W)","Measured Current Input (A)","Notes: in operating mode, H = heating, C = cooling, D = dehumidifier")]
    for row in table_test.rows:     #set the row height 
        row.height=Pt(20)
    column_cells=table_test.columns[0].cells  #choose the first column of table
    for t,u,v,w,x,y,z,s in values:
        column_cells[0].text=t
        column_cells[1].text=u
        column_cells[2].text=v
        column_cells[3].text=w
        column_cells[4].text=x
        column_cells[5].text=y
        column_cells[6].text=z
        column_cells[7].text=s
    table_test.cell(7,0).merge(table_test.cell(7,4))      ##merge the cells
    document.add_page_break()
#    table_test=document.add_table(rows=1,cols=1,style="Table Grid")

def Clause11():
    p=document.add_paragraph().add_run("Test Item:\tHeating test(Clause 11.8)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\t Pass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_couple=document.add_table(rows=25, cols=4, style="Table Grid")        ##Table for thermocouple
    row_cells=T_couple.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    col_cells=T_couple.columns[0].cells     #choose the first column of T_couple table
    for i in range(1,25):   #set the item no series 
        col_cells[i].text=str(i)
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)

    document.add_paragraph()        ##blank line
    T1_pressure=document.add_table(rows=1, cols=1, style="Table Grid")        ##Table1 for pressure 
    T1_pressure.rows[0].cells[0].text="Record the Max. working pressure on both high/low side"
    T2_pressure=document.add_table(rows=2, cols=2, style="Table Grid")        ##Table2 for pressure 
    col1=T2_pressure.columns[0].cells       ##select the first column of Table2 for pressure
    col1[0].text="Condenser side (MPa)"
    col1[1].text="Evaporator side (MPa)"
    for row in T1_pressure.rows:     #set the row height
        row.height=Pt(20)
    for row in T2_pressure.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

    p=document.add_paragraph().add_run("Test Item:\tMeasurement of winding temperaturer rise(Clause 11.8)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    T_head=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    T_head.rows[1].height=Pt(20)       ##set the row height
    col1=T_head.columns[0].cells    #first column
    col2=T_head.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):\tt1=\n\t\t\tt2="
    col2[1].text="Test frequency(Hz):"

    T_resistance=document.add_table(rows=4, cols=5, style="Table Grid")        ##Table for resistance method
    row_cells=T_resistance.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Temperature rise of winding"
    row_cells[1].text="R1(ohm)"
    row_cells[2].text="R2(ohm)"
    row_cells[3].text="dT(C)"
    row_cells[4].text="Limitation(C)"
    for row in T_resistance.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause13():
    p=document.add_paragraph().add_run("Test Item:\tLeakage current test & Electric strength under normal operation (Clause 13.2/13.3)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    p=document.add_paragraph().add_run("Clause 13.2").font.name="Arial"
    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_leakage=document.add_table(rows=3, cols=3, style="Table Grid")    ##Table for leakage current test  
    col1=T_leakage.columns[0].cells  ##select the first column 
    col3=T_leakage.columns[2].cells  ##select the third column 
    col1[0].text="L/N - Earthing metal part"
    col1[2].text="L/N - Enclosure (with metal foil or unground metal part"
    col3[0].text="0.75mA for portable appliance"
    col3[1].text="(    )mA for stationary appliance"
    col3[2].text="0.35mA peak"
    ##table format
    T_leakage.rows[0].height=Pt(20)     ##set the height of first row
    T_leakage.rows[1].height=Pt(20)     ##set the height of first row
    T_leakage.cell(0,0).merge(T_leakage.cell(1,0))      ##merge the cells

    document.add_paragraph()        ##blank line
    p=document.add_paragraph().add_run("Clause 13.3").font.name="Arial"
    T_dielectric=document.add_table(rows=6, cols=3, style="Table Grid")    ##Table for dielectric strength test  
    row1=T_dielectric.rows[0].cells  ##select the first row 
    row1[0].text="Test voltage applied between:"
    row1[1].text="Test voltage(V):"
    row1[2].text="Brokedown:"
    row2=T_dielectric.rows[1].cells  ##select the second row 
    row2[0].text="SELV isolated with basic insulation"
    row2[1].text="500"
    row2[2].text="Yes/No"
    row3=T_dielectric.rows[2].cells  ##select the third row 
    row3[0].text="Between live part and the earthing metal enclosure (basic insulation)"
    row3[1].text="1000"
    row3[2].text="Yes/No"
    row4=T_dielectric.rows[3].cells  ##select the fourth row 
    row4[0].text="Between basic insulation part and the non-metal enclosure (supplementary insulation)"
    row4[1].text="1750"
    row4[2].text="Yes/No"
    row5=T_dielectric.rows[4].cells  ##select the fifth row 
    row5[0].text="Between live part and non-metal enclosure or SELV (reinforce insulation)"
    row5[1].text="3000"
    row5[2].text="Yes/No"
    row6=T_dielectric.rows[5].cells  ##select the sixth row 
    row6[0].text="Note:\n1: Metal foil having an area not exceeding 20 cm x 10 cm which is in contact with accessible surfaces of insulating materials.\n2: Protective impedance and radio interference filters are disconnected before carrying out the tests."
    ##table format
    for row in T_dielectric.rows:     #set the row height
        row.height=Pt(20)
    T_dielectric.cell(5,0).merge(T_dielectric.cell(5,2))      ##merge the cells
    document.add_page_break()

def Clause16():
    p=document.add_paragraph().add_run("Test Item:\tLeakage current test & Electric strength after humidity (Clause 16.2/16.3)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    p=document.add_paragraph().add_run("Clause 16.2").font.name="Arial"
    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Humidity(%RH):"
    T_leakage=document.add_table(rows=3, cols=3, style="Table Grid")    ##Table for leakage current test  
    col1=T_leakage.columns[0].cells  ##select the first column 
    col3=T_leakage.columns[2].cells  ##select the third column 
    col1[0].text="L/N - Earthing metal part"
    col1[2].text="L/N - Enclosure (with metal foil or unground metal part"
    col3[0].text="0.75mA for portable appliance"
    col3[1].text="(    )mA for stationary appliance"
    col3[2].text="0.25mA"
    ##table format
    T_leakage.rows[0].height=Pt(20)     ##set the height of first row
    T_leakage.rows[1].height=Pt(20)     ##set the height of first row
    T_leakage.cell(0,0).merge(T_leakage.cell(1,0))      ##merge the cells

    document.add_paragraph()        ##blank line
    p=document.add_paragraph().add_run("Clause 16.3").font.name="Arial"
    T_dielectric=document.add_table(rows=6, cols=3, style="Table Grid")    ##Table for dielectric strength test  
    row1=T_dielectric.rows[0].cells  ##select the first row 
    row1[0].text="Test voltage applied between:"
    row1[1].text="Test voltage(V):"
    row1[2].text="Brokedown:"
    row2=T_dielectric.rows[1].cells  ##select the second row 
    row2[0].text="SELV isolated with basic insulation"
    row2[1].text="500"
    row2[2].text="Yes/No"
    row3=T_dielectric.rows[2].cells  ##select the third row 
    row3[0].text="Between live part and the earthing metal enclosure (basic insulation)"
    row3[1].text="1250"
    row3[2].text="Yes/No"
    row4=T_dielectric.rows[3].cells  ##select the fourth row 
    row4[0].text="Between basic insulation part and the non-metal enclosure (supplementary insulation)"
    row4[1].text="1750"
    row4[2].text="Yes/No"
    row5=T_dielectric.rows[4].cells  ##select the fifth row 
    row5[0].text="Between live part and non-metal enclosure or SELV (reinforce insulation)"
    row5[1].text="3000"
    row5[2].text="Yes/No"
    row6=T_dielectric.rows[5].cells  ##select the sixth row 
    row6[0].text="Note:\n1: Metal foil having an area not exceeding 20 cm x 10 cm which is in contact with accessible surfaces of insulating materials.\n2: Protective impedance and radio interference filters are disconnected before carrying out the tests."
    ##table format
    for row in T_dielectric.rows:     #set the row height
        row.height=Pt(20)
    T_dielectric.cell(5,0).merge(T_dielectric.cell(5,2))      ##merge the cells
    document.add_page_break()

def Clause17():
    p=document.add_paragraph().add_run("Test Item:\tOverload of transformer test(Clause 17)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_couple=document.add_table(rows=5, cols=4, style="Table Grid")        ##Table for thermocouple
    row_cells=T_couple.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    col_cells=T_couple.columns[0].cells     #choose the first column of T_couple table
    for i in range(1,5):   #set the item no series 
        col_cells[i].text=str(i)
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)

    T_resistance=document.add_table(rows=3, cols=6, style="Table Grid")        ##Table for resistance method
    row_cells=T_resistance.rows[0].cells  #choose the first row of T_resistance table
    row_cells[0].text="Temperature rise of winding"
    row_cells[1].text="R1(ohm)"
    row_cells[2].text="R2(ohm)"
    row_cells[3].text="dT(C)"
    row_cells[4].text="Limitation(C)"
    row_cells[5].text="Insulation class"
    row_cells=T_resistance.rows[1].cells  #choose the second row of T_resistance table
    row_cells[0].text="Primary winding"
    row_cells[1].text="\n(t1=   C)"
    row_cells[2].text="\n(t2=   C)"
    row_cells=T_resistance.rows[2].cells  #choose the third row of T_resistance table
    row_cells[0].text="secondary winding"
    row_cells[1].text="\n(t1=   C)"
    row_cells[2].text="\n(t2=   C)"
    document.add_page_break()

def Clause19_2and19_3():
    ##clause 19.2&19.3
    p=document.add_paragraph().add_run("Test Item:\tAbnormal Operation Restricted Heat Dissipation & Overload Test (clause 19.2&19.3)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_couple=document.add_table(rows=26, cols=4, style="Table Grid")        ##Table for thermocouple
    row_cells=T_couple.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    col_cells=T_couple.columns[0].cells     #choose the first column of T_couple table
    col_cells[25].text="When the protective device operates, the temperature of air outlet is ______ C."
    T_couple.cell(25,0).merge(T_couple.cell(25,3))      ##merge the cells
    for i in range(1,25):   #set the item no series 
        col_cells[i].text=str(i)
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)

    document.add_paragraph()        ##blank line
    T1_pressure=document.add_table(rows=1, cols=1, style="Table Grid")        ##Table1 for pressure 
    T1_pressure.rows[0].cells[0].text="Record the Max. working pressure on both high/low side"
    T2_pressure=document.add_table(rows=2, cols=2, style="Table Grid")        ##Table2 for pressure 
    col1=T2_pressure.columns[0].cells       ##select the first column of Table2 for pressure
    col1[0].text="Condenser side (MPa)"
    col1[1].text="Evaporator side (MPa)"
    for row in T1_pressure.rows:     #set the row height
        row.height=Pt(20)
    for row in T2_pressure.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()
    
def Clause19_7():
    ##clause 19.7
    p=document.add_paragraph().add_run("Test Item:\tLocking test for the motor (clause 19.7)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_couple=document.add_table(rows=5, cols=4, style="Table Grid")        ##Table for thermocouple
    T_couple.rows[0].cells[0].text="Test duration:\n15 days (360 h) or protection device permanently opens the circuit"     ##fill in words in first line 
    T_couple.cell(0,0).merge(T_couple.cell(0,3))      ##merge the cells
    row_cells=T_couple.rows[1].cells  #choose the second row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    row_cells=T_couple.rows[2].cells  #choose the third row of T_couple table
    row_cells[0].text="1"
    row_cells[1].text="Motor winding"
    row_cells=T_couple.rows[3].cells  #choose the fourth row of T_couple table
    row_cells[0].text="2"
    row_cells[1].text="Enclosure"
    T_couple.rows[4].cells[0].text="Notes:\nThree days (72 h) after the beginning of the test, the motor shall withstand an electric strength test as specified in 16.3.\nAt the end of the test, the leakage current test is applied (twice the rated voltage) between all windings and the enclosure, the value is _____mA, do not exceed 2mA.\n"     ##fill in words in fifth line 
    T_couple.cell(4,0).merge(T_couple.cell(4,3))      ##merge the cells
    row_cells=T_couple.rows[1].cells  #choose the second row of T_couple table
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause19_11_2():
    ##Clause 19.11.2
    p=document.add_paragraph().add_run("Test Item:\tFault condition of electric circut (clause 19.11.2)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_fault=document.add_table(rows=8,cols=1,style="Table Grid")     ##The table for fault condition
    col1=T_fault.columns[0].cells    #first column
    col1[0].text="-Short circuit of creepage and clearance between live parts of different potential, if these distances are less than the values specified in 29.1. Following parts are short-circuited:"
    col1[1].text="-Open circuit at the terminals of any component. Following components are opened:"
    col1[2].text="-Short circuit of capacitors not complying with IEC 384-14. Following capacitors are short-circuited:"
    col1[3].text="-Short circuit of any two terminals of an electronic component, other than integrated circuits. This fault condition is not applied between the two circuits of an optocoupler. Following electronic components are short-circuited:"
    col1[4].text="-Failure of triacs in the diode mode. Following triacs are in diode mode:"
    col1[5].text="-Failure of an integrated circuit. Following ICs are disabled:"
    col1[6].text="-Failure of an electronic power switching device in a partial turn-on mode with loss of gate (base) control. During this test, winding temperatures shall not exceed the values given in 19.7.:"
    col1[7].text="Notes:\nNo hazard happens in these situations and all results comply with 19.13."
    for row in T_fault.rows:     #set the row height
        row.height=Pt(40)
    document.add_page_break()
   
def Clause19_11_4_8():
    ##Clause 19.11.4.8
    p=document.add_paragraph().add_run("Test Item:\tVoltage drop test (clause 19.11.4.8)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    T_voltage=document.add_table(rows=4, cols=2, style="Table Grid")        ##Table for voltage drop
    T_voltage.rows[0].cells[0].text="Record voltage: ______ V\n(The voltage that the appliance ceases to respond to user inputs or parts controlled by the programmable component cease to operate, whichever occurs first.)"     ##fill in words in first line 
    T_voltage.cell(0,0).merge(T_voltage.cell(0,1))      ##merge the cells
    row_cells=T_voltage.rows[1].cells  #choose the second row of T_voltage table
    row_cells[0].text="The condition of appliance after the test:"
    row_cells=T_voltage.rows[2].cells  #choose the third row of T_voltage table
    row_cells[0].text="- Continue operating normally"
    row_cells[1].text="Yes / No"
    row_cells=T_voltage.rows[3].cells  #choose the fourth row of T_voltage table
    row_cells[0].text="- manual operation shall be required to restart it."
    row_cells[1].text="Yes / No"
    document.add_page_break()

def Clause19_101():
    ##Clause 19.101
    p=document.add_paragraph().add_run("Test Item:\tRestriction of Heat Transfer Medium Flow (clause 19.101)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"

    T_couple=document.add_table(rows=26, cols=4, style="Table Grid")        ##Table for thermocouple
    row_cells=T_couple.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    col_cells=T_couple.columns[0].cells     #choose the first column of T_couple table
    T_couple.cell(25,0).merge(T_couple.cell(25,3))      ##merge the cells
    T_couple.rows[25].cells[0].text="Notes:\n-The heat transfer medium flow of the outdoor heat exchanger is restricted or shut off\n-The heat transfer medium flow, fluid or air, of the indoor heat exchanger, restricted or shut off\n-Appliances incorporating a motor common to both the indoor and outdoor heat exchangers are subjected to the above test the motor being disconnected"
    for i in range(1,25):   #set the item no series 
        col_cells[i].text=str(i)
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)

    document.add_paragraph()        ##blank line
    T1_pressure=document.add_table(rows=1, cols=1, style="Table Grid")        ##Table1 for pressure 
    T1_pressure.rows[0].cells[0].text="Record the Max. working pressure on both high/low side"
    T2_pressure=document.add_table(rows=2, cols=2, style="Table Grid")        ##Table2 for pressure 
    col1=T2_pressure.columns[0].cells       ##select the first column of Table2 for pressure
    col1[0].text="Condenser side (MPa)"
    col1[1].text="Evaporator side (MPa)"
    for row in T1_pressure.rows:     #set the row height
        row.height=Pt(20)
    for row in T2_pressure.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause19_102():
    ##Clause 19.102
    p=document.add_paragraph().add_run("Test Item:\tAbnormal Temperature of Indoor Water (clause 19.102)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"

    T_couple=document.add_table(rows=26, cols=4, style="Table Grid")        ##Table for thermocouple
    row_cells=T_couple.rows[0].cells  #choose the first row of T_couple table
    row_cells[0].text="Item No."
    row_cells[1].text="Thermocouple location"
    row_cells[2].text="Actual temperature(C)"
    row_cells[3].text="Limitation(C)"
    col_cells=T_couple.columns[0].cells     #choose the first column of T_couple table
    T_couple.cell(25,0).merge(T_couple.cell(25,3))      ##merge the cells
    T_couple.rows[25].cells[0].text="Notes:\nThe indoor water temperature shall be raised 15 K with a rate of 2 K/min and this temperature maintained for 30 min, after which the water temperature is lowered to its original value at the same velocity."
    for i in range(1,25):   #set the item no series 
        col_cells[i].text=str(i)
    for row in T_couple.rows:     #set the row height
        row.height=Pt(20)

    document.add_paragraph()        ##blank line
    T1_pressure=document.add_table(rows=1, cols=1, style="Table Grid")        ##Table1 for pressure 
    T1_pressure.rows[0].cells[0].text="Record the Max. working pressure on both high/low side"
    T2_pressure=document.add_table(rows=2, cols=2, style="Table Grid")        ##Table2 for pressure 
    col1=T2_pressure.columns[0].cells       ##select the first column of Table2 for pressure
    col1[0].text="Condenser side (MPa)"
    col1[1].text="Evaporator side (MPa)"
    for row in T1_pressure.rows:     #set the row height
        row.height=Pt(20)
    for row in T2_pressure.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause19_103():
    ##Clause 19.103
    p=document.add_paragraph().add_run("Test Item:\tAbnormal Ambient Temperature (clause 19.103)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=2,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"

    T_result=document.add_table(rows=1,cols=1,style="Table Grid")   ##The result of table
    T_result.rows[0].cells[0].text="1. dry-bulb temperature is increased to a value 10 K above the maximum temperature specified by the manufacturer:_____C\n2. dry-bulb temperature is reduced to a value 5 K below the maximum temperature specified by the manufacturer:_____C\nAfter testing, no hazard situation was occured."
    document.add_page_break()

def Clause19_104():
    ##Clause 19.104
    p=document.add_paragraph().add_run("Test Item:\tCover Test for Appliance with Supplementary Heaters (clause 19.104)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    table_test=document.add_table(rows=3,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    col1=table_test.columns[0].cells    #first column
    col2=table_test.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    table_test.cell(2,0).merge(table_test.cell(2,1))        #merge the cells
    table_test.rows[2].cells[0].text="Result:\na)No emit flames, molten metal, poisonous or ignitable gas in hazard amount.\nb)The temperature do/do not exceed 150C.\nc)Thermal protective devices operated during the test."
    document.add_page_break()

def AnnexEE():
    p=document.add_paragraph().add_run("Test Item:\tPressure Tests (Annex EE)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"
    
    table_test=document.add_table(rows=3,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    row1=table_test.rows[0].cells    #first rows 
    row2=table_test.rows[1].cells    #second rows 
    row3=table_test.rows[2].cells    #third rows 
    table_test.cell(1,0).merge(table_test.cell(1,1))        #merge the second line
    table_test.cell(2,0).merge(table_test.cell(2,1))        #merge the third line
    row1[0].text="Model:"
    row1[1].text="Ambient(C):"
    row2[0].text="The test pressure is the maximum of the following:\na)Three times the maximum allowable pressure developed during operation under Clause 11.\n\tHigh side: [\t] MPa, low side: [\t] MPa\n\nb)Three times the maximum allowable pressure developed during abnormal operation under Clause 19.\n\tHigh side: [\t] MPa, low side: [\t] MPa\n\nc)Three times the maximum allowable pressure developed during standstill. In order to determine the standstill pressure, the appliance shall be soaked in the highest operating temperature specified by the manufacturer for 1 h with power off.\n\tHigh side: [\t] MPa, low side: [\t] MPa\n\nThe test pressure applied: High side: [\t] MPa, low side: [\t] MPa\nThe pressure is raised gradually until the required test pressure is reached.\nThe pressure is maintained for at least 1 min, during which time the sample shall not leak.\n"
    row3[0].text="Notes:\na)Pressure gauges and control mechanisms need not be subjected to the test, provided the parts meet the requirements of the component.\nb)Where gaskets are employed for sealing parts under pressure, leakage at gaskets is acceptable, provided the leakage only occurs at a value greater than 120% of the maximum allowable pressure.\n"
    document.add_page_break()

def AnnexT():
    p=document.add_paragraph().add_run("Test Item:\tUV-C Radiation (Annex T)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"
    
    document.add_paragraph().add_run("Table T.1-Minimum property retention limits after UV-C exposure").font.name="Arial"
    table_T1=document.add_table(rows=7,cols=7,style="Table Grid")     ##Table T.1-Minimum property retention limits after UV-C exposure
    row1=table_T1.rows[0].cells    #first rows 
    row2=table_T1.rows[1].cells    #second rows 
    row5=table_T1.rows[4].cells    #fifth rows 
    row1[0].text="Parts to be Tested"
    row1[1].text="Test Item"
    row1[2].text="Tensile Strength"
    row1[3].text="Flexural strength"
    row1[4].text="Charpy impact"
    row1[5].text="Izod impact"
    row1[6].text="Minimum retention after testing"
    row2[0].text="Parts providing mechanical support"
    row5[0].text="Parts providing impact resistance"
    table_T1.cell(1,0).merge(table_T1.cell(3,0))      ##merge the cells
    table_T1.cell(4,0).merge(table_T1.cell(6,0))      ##merge the cells
    document.add_paragraph()        ##blank line
    ##table format
    for row in table_T1.rows:     #set the row height
        row.height=Pt(20)
    
    document.add_paragraph().add_run("Table T.2-Minimum electric strength for internal wiring after UV-C exposure").font.name="Arial"
    table_T2=document.add_table(rows=4,cols=4,style="Table Grid")     ##Table T.2-Minimum electric strength for internal wiring after UV-C exposure
    row1=table_T2.rows[0].cells    #first rows 
    row2=table_T2.rows[1].cells    #second rows 
    row1[0].text="Parts to be Tested"
    row1[1].text="Test Item"
    row1[2].text="Color"
    row1[3].text="Compliance"
    row2[0].text="Electrical insulation of internal wiring"
    table_T2.cell(1,0).merge(table_T2.cell(3,0))      ##merge the cells
    ##table format
    for row in table_T2.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause22_104():
    ##Clause 22.104
    p=document.add_paragraph().add_run("Test Item:\tWater Pressure Test for Containers (Clause 22.104)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"
    
    table_test=document.add_table(rows=3,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in table_test.rows:     #set the row height
        row.height=Pt(20)
    row1=table_test.rows[0].cells    #first rows 
    row2=table_test.rows[1].cells    #second rows 
    row3=table_test.rows[2].cells    #third rows 
    table_test.cell(1,0).merge(table_test.cell(1,1))        #merge the second line
    table_test.cell(2,0).merge(table_test.cell(2,1))        #merge the third line
    row1[0].text="Model:"
    row1[1].text="Ambient(C):"
    row2[0].text="The test pressure is [\t] MPa, applied on _____. The test pressure is determined as following:\n\na) twice the permissible excessive operating pressure for closed containers.\n\nb) 0.15 MPa for open containers.\n\nAfter the test, no water have leaked out and the containers have not ruptured.\n"
    row3[0].text="Notes:\nWater pressure is raised at a rate 0.13 MPa per seconed and maintained at that value for 5 minutes.\n"
    document.add_page_break()

def Clause30_1():
    ##Clause 30.1
    p=document.add_paragraph().add_run("Test Item:\tBall Pressure Test (Clause 30.1)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    T_ball=document.add_table(rows=1,cols=3,style="Table Grid")     ##The head of table
    T_ball.rows[0].cells[0].text="Model:"
    T_ball.rows[0].cells[1].text="Ambient(C):"
    T_ball.rows[0].cells[2].text="Humidity(%RH):"
    T_data=document.add_table(rows=7,cols=6,style="Table Grid")     ##The data of table
    row1=T_data.rows[0].cells   #first row of data table
    col1=T_data.columns[0].cells #first column of data table
    row1[0].text="Item"
    row1[1].text="Parts"
    row1[2].text="Thickness\n(mm)"
    row1[3].text="Test Temperature\n(C)"
    row1[4].text="Impression Diameter\n(mm)"
    row1[5].text="Result"
    for i in range(1,7):
        col1[i].text=str(i)
    ##table format
    for row in T_ball.rows:     #set the row height
        row.height=Pt(20)
    for row in T_data.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def Clause30_2():
    ##Clause 30.2
    p=document.add_paragraph().add_run("Test Item:\tGlow Wire Test & Needle Flame Test (Clause 30.2.3 & 30.2.4)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    T_data=document.add_table(rows=9,cols=9,style="Table Grid")     ##The data of table
    T_data.cell(0,0).merge(T_data.cell(0,2))  #merge the cells
    T_data.cell(0,3).merge(T_data.cell(0,8))  #merge the cells
    T_data.cell(1,0).merge(T_data.cell(2,0))  #merge the cells
    T_data.cell(1,1).merge(T_data.cell(1,2))  #merge the cells
    T_data.cell(1,3).merge(T_data.cell(1,6))  #merge the cells
    T_data.cell(1,7).merge(T_data.cell(2,7))  #merge the cells
    T_data.cell(1,8).merge(T_data.cell(2,8))  #merge the cells
    T_data.rows[0].cells[0].text="Model:"
    T_data.rows[0].cells[3].text="Ambient(C):"
    T_data.rows[1].cells[0].text="Parts"
    T_data.rows[1].cells[1].text="Tracking Test (V)"
    T_data.rows[2].cells[1].text="175"
    T_data.rows[2].cells[2].text="250"
    T_data.rows[1].cells[3].text="Grow Wiring Test (C)"
    T_data.rows[2].cells[3].text="550"
    T_data.rows[2].cells[4].text="650"
    T_data.rows[2].cells[5].text="750"
    T_data.rows[2].cells[6].text="850"
    T_data.rows[1].cells[7].text="Needle Flame"
    T_data.rows[1].cells[8].text="Result"
    for row in T_data.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def AnnexN():
    p=document.add_paragraph().add_run("Test Item:\tProof Tracking Test (Annex N)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"
    
    T_PTI=document.add_table(rows=5,cols=4,style="Table Grid")     ##The table for PTI
    T_sample=document.add_table(rows=5,cols=7,style="Table Grid")     ##The table for sample 
    row1=T_PTI.rows[0].cells  #first row
    col1=T_PTI.columns[0].cells #first column
    row1s=T_sample.rows[0].cells  #first row
    col1s=T_sample.columns[0].cells #first column
    row1[0].text="Item"
    row1[1].text="Parts"
    row1[2].text="Thickness\n(mm)"
    row1[3].text="PTI\n(v)"
    row1s[0].text="Item"
    row1s[6].text="Result"
    for i in range(1,5):
        col1[i].text=str(i)
    for i in range(1,6):
        row1s[i].text=str(i)+"#"
    for i in range(1,5):
        col1s[i].text=str(i)
    for row in T_PTI.rows:     #set the row height
        row.height=Pt(20)
    for row in T_sample.rows:     #set the row height
        row.height=Pt(20)
    document.add_page_break()

def AnnexFF():
    p=document.add_paragraph().add_run("Test Item:\tLeakage Simulation Tests (Annex FF)").font.name="Arial"
    p=document.add_paragraph().add_run("Test Method:\tAs specified in standard").font.name="Arial"
    p=document.add_paragraph().add_run("Test Result:\tPass/Failed").font.name="Arial"

    T_head=document.add_table(rows=4,cols=2,style="Table Grid")     ##The head of table
    ##table format
    for row in T_head.rows:     #set the row height
        row.height=Pt(20)
    col1=T_head.columns[0].cells    #first column
    col2=T_head.columns[1].cells    #second column
    col1[0].text="Model:"
    col1[1].text="Test voltage(V):"
    col2[0].text="Test condition(C):"
    col2[1].text="Test frequency(Hz):"
    col1[2].text="Refrigerant:\n\nRefrigerant charge (g):\n\nLFL:\n\nThe minimum volume: V=(15xMc)/LFL=\t\tm3\n\nMaximum concentration measured:\n"
    col1[3].text="Result:\n\nThe maximum concentration measured do not exceed 25% of LFL of refrigerant.\tYes/No\n\nThe average concentration measured do not exceed 15% of LFL of refrigerant during the test.\tYes/No\n\n"
    T_head.cell(2,0).merge(T_head.cell(2,1))    #merge the cells
    T_head.cell(3,0).merge(T_head.cell(3,1))    #merge the cells

def Check(selection):
###function to detect which testing clause is choosed.
    if selection=="1":
        Clause10()
    elif selection=="2":
        Clause11()
    elif selection=="3":
        Clause13()
    elif selection=="4":
        Clause16()
    elif selection=="5":
        Clause17()
    elif selection=="6":
        Clause19_2and19_3()
    elif selection=="7":
        Clause19_7()
    elif selection=="8":
        Clause19_11_2()
    elif selection=="9":
        Clause19_11_4_8()
    elif selection=="10":
        Clause19_101()
    elif selection=="11":
        Clause19_102()
    elif selection=="12":
        Clause19_103()
    elif selection=="13":
        Clause19_104()
    elif selection=="14":
        AnnexEE()
    elif selection=="15":
        AnnexT()
    elif selection=="16":
        Clause22_104()
    elif selection=="17":
        Clause30_1()
    elif selection=="18":
        Clause30_2()
    elif selection=="19":
        AnnexN()
    elif selection=="20":
        AnnexFF()
    elif selection=="1.1":
        Clause10()
        Clause11()
        Clause13()
        Clause16()
        Clause19_101()
        Clause19_103()
    elif selection=="1.2":
        Clause10()
        Clause11()
        Clause13()
        Clause16()
        Clause19_7()
        Clause19_101()
        Clause19_103()
    elif selection=="1.3":
        Clause10()
        Clause11()
        Clause13()
        Clause16()


###对选择进行排序然后输出测试内容
def selection_sort(selection):
    '''
    selection(str):测试的选择
    '''
    #对选择字符串做处理
    selection_sort=[]
    prehandle=selection.split(",")
    for i in prehandle:
        if "-" in i:
            j=i.split("-")
            j[0]=int(j[0])
            j[1]=int(j[1])
            for k in range(j[0],j[1]+1):
                k=str(k)
                selection_sort.append(k)
        else:
            selection_sort.append(i)
    #转化为整数
    selection=list(map(int,selection_sort))
    #对整数排序
    selection.sort()
    #转化为字符串
    selection=list(map(str,selection))
    #遍历每一个选择并输出测试章节
    for value in selection:
        Check(value)


def Menu(Test_clauses):
    print("1.Regular combination")
    print("2.Customize")
    while True:
        sel=input("Please choose:")
        if sel=="1":#固定的测试组合
            print("1.1 alternative certified compressor or motor")
            print("1.2 alternative uncertified compressor or motor")
            print("1.3 alternative uncerticied transformer")
            while True:                
                sel=input("Please choose:")
                if (sel=="1.1" or sel=="1.2") or sel=="1.3":
                    Check(sel)
                    break   #jump out the second loop
                else:
                    print("WARNING: wrong input, please choose again!:")
            break   #jump out the first loop
        elif sel=="2":#自定义测试内容
            n=1
            for i in Test_clauses:
                print(str(n)+" "+i[1]+"==>"+i[0])
                n=n+1
            selection=input("Please select the testing clauses:")
            selection_sort(selection)#对选择的测试范围进行处理
            break   #jump out the first loop
        else:
            print("WARNING: wrong input, please choose again!:")

####################################### main program #############################################
if __name__=='__main__':
    while True:
        try:
            print("This TDS is for IEC 60335-2-40:2018 in conjunction with IEC 60335-1:2010+A1:2013+A2:2016!")
            print("================================================== Program begin ==================================================")
            job=input("Please input the project number:")
            print('project name:',job)
            if os.name=='nt':
                document = Document('.\Temp.docx') #Open the template document
            elif os.name=='posix':
                document = Document('./Temp.docx') #Open the template document
            print('opening the document')
            body=document.add_paragraph()
            Content(values)
    
            #选择测试内容并输出
            Menu(Test_clauses)
            document.save(job+'.docx')
    
            print("==================================================Program END ==================================================") 
            flag=input("Press Enter to continue! Others to EXIT!")
            if flag!="":
                break
            else:
                pass
    #            os.system("cls")
    
        except:
            print("==================================================Program END ==================================================") 
            print("Error, please contact Michael.jc.jin@intertek.com")
            break

